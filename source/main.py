import os
from docx import Document

filename = ''
folder_name = 'CODE'
folder_name_len = len(folder_name) + 3
sample_file = 'template.docx'
type_file = []
prodl_name = ""
with open("..\\settings.txt", "r",  encoding='utf-8', errors='ignore') as f:
    tmp = f.read().split("\n")
    type_file = tmp[0].replace(" ", "").split(",")
    filename = "..\\" + str(tmp[1]) + " 12 01 (Текст программы)"
    prodl_name = tmp[2]

def check(file_name):
    if file_name.split(".")[-1] in type_file:
        return True
    return False

def read_file(filename):
    with open(filename, "r",  encoding='utf-8', errors='ignore') as f:
        return f.read()

folder = []
for i in os.walk("..\\" + folder_name):
    folder.append(i)

paths = []
for address, dirs, files in folder:
    for file in files:
        if check(file):
            paths.append(address + '\\' + file)

total_code_file = []
print(paths)
for i, path in enumerate(paths):
    if i == 0:
        catalog = path.split('\\')[1]
        # total_code_file.append('Каталог ' + catalog + '\n')
    if path[folder_name_len:].split('\\')[1] != catalog:
        catalog = path[folder_name_len:].split('\\')[1]
        total_code_file.append('Каталог ' + catalog + '\n')

    total_code_file.append('Файл ' + path[folder_name_len:] + '\n')
    total_code_file.append(read_file(path))
    total_code_file += '\n'

doc = Document(sample_file)
for p in doc.paragraphs:
    if '<КОД ПРОГРАММЫ>' in p.text:
        p.text = ''
        for line in total_code_file:
            if line.rstrip() > '' and line.split()[0] == 'Каталог':
                p.insert_paragraph_before(line.rstrip(), 'Heading 2')
            elif line.rstrip() > '' and line.split()[0] == 'Файл':
                p.insert_paragraph_before(line.rstrip(), 'Heading 3')
            else:
                p.insert_paragraph_before(line.rstrip(), 'КОД')
    elif "<НАЗВАНИЕ>" in p.text:
        p.text = ''
        p.insert_paragraph_before(prodl_name, "Title")
    elif "<номер>" in p.text:
        p.text = ''
        p.insert_paragraph_before(filename[3:-18], "Title")
    elif "<номе2р>" in p.text:
        p.text = ''
        p.insert_paragraph_before(filename[3:-18] + "-ЛУ", "Title")

    elif "<номе1р>" in p.text:
        p.text = ''
        p.insert_paragraph_before("             " + filename[3:-18] + "-ЛУ", "еспд-дец-1")

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if "643.ХХХХ.ХХХХХ-01 12 01" in paragraph.text:
                paragraph.text = ''
                paragraph.insert_paragraph_before(filename[3:-18], "Title")

doc.save(filename + '.docx')
